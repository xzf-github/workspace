from docx import Document
doc = Document('temp.docx')
contents = ''.join((p.text for p in doc.paragraphs))
words = []
#enumerate函数用于获取字符串中每个字符的索引和值。
#contents[:-2]`表示从字符串的开始到倒数第二个字符。
#循环遍历每个字符及其索引。
for index, ch in enumerate(contents[:-2]):
    #检查当前字符是否与其后两个字符相同
    if ch==contents[index+1] or ch==contents[index+2]:
        word = contents[index:index+3]
        #如果满足条件，提取这三个字符并将其添加到列表中（如果未在列表中）
        if word not in words:
            words.append(word)
            print(word)

'''
from docx import Document  
  
doc = Document(doc_path)  
for para in doc.paragraphs:  
    text = para.text  
    for i in range(len(text) - min_length + 1):  
        if text[i] == text[i + min_length - 1]:  
            print(f"连续重复字在段落 '{para.text}' 中从索引 {i} 开始")  
doc.close()  


'''